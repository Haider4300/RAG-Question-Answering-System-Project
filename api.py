from dotenv import load_dotenv
load_dotenv()
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from contextlib import asynccontextmanager
from typing import Optional, List
from datetime import datetime
import asyncio
import logging
import threading
import os
import uuid
import json
import json as _json
import time

from langchain_community.tools import DuckDuckGoSearchRun
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ── Cloud LLM + Embeddings (replaces Ollama for HF Spaces deployment) ─────────
# Groq: free API, ~1s responses, same LangChain interface as ChatOllama
# HuggingFace embeddings: runs locally inside the container, no API key needed
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Constants ────────────────────────────────────────────────────────────────
STORAGE_DIR          = "uploaded_files"
CHAT_HISTORY_FILE    = "chat_history.json"
DOCUMENTS_META_FILE  = "documents_meta.json"
INDEX_DIR            = "faiss_index_uploaded"
CHUNK_SIZE           = 1000
CHUNK_OVERLAP        = 200
TOP_K                = 3
TOP_K_BROAD          = 6
MAX_CONTEXT_CHARS    = 3000
EMBEDDING_BATCH_SIZE = 20
MAX_FILE_SIZE        = 10 * 1024 * 1024
MAX_PAGES            = 250
MAX_CHUNKS           = 500
RELEVANCE_THRESHOLD  = 0.40

# ─── Globals ──────────────────────────────────────────────────────────────────
_llm         = None
_embeddings  = None
_vectorstore = None
_retriever   = None
_web_search  = None

_processing_lock = threading.Lock()

documents_db: dict = {}
chats_db:     dict = {}
_active_doc:  dict = {}
doc_status:   dict = {}

os.makedirs(STORAGE_DIR, exist_ok=True)
os.makedirs(INDEX_DIR,   exist_ok=True)


# ─── Lazy singletons ──────────────────────────────────────────────────────────

def get_llm() -> ChatGroq:
    global _llm
    if _llm is None:
        _llm = ChatGroq(
           model="llama-3.1-8b-instant",
            temperature=0,
            max_tokens=512,
            api_key=os.getenv("GROQ_API_KEY"),
        )
    return _llm


def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embeddings


def get_web_search() -> DuckDuckGoSearchRun:
    global _web_search
    if _web_search is None:
        _web_search = DuckDuckGoSearchRun()
    return _web_search


# ─── Query classifier ─────────────────────────────────────────────────────────

def classify_query(question: str) -> str:
    q = question.strip().lower()

    GREETING_EXACT = {
        "hi", "hey", "hello", "hiya", "howdy", "greetings",
        "sup", "yo", "heya", "hai", "bye", "goodbye", "thanks",
        "thank you", "ok", "okay", "sure", "great", "nice",
    }
    GREETING_STARTERS = (
        "hi ", "hi!", "hi,", "hey ", "hey!", "hey,",
        "hello ", "hello!", "hello,", "hiya", "howdy",
        "good morning", "good afternoon", "good evening", "good night",
        "what's up", "whats up", "how are you", "how r u",
    )
    SMALL_TALK_EXACT = {
        "what's going on", "whats going on", "what is going on",
        "what's happening", "whats happening", "what's new", "whats new",
        "how's it going", "hows it going", "what can you do", "what do you do",
        "who are you", "what are you", "who made you",
        "what's your name", "whats your name", "what is your name",
        "are you an ai", "are you a bot", "are you real",
        "tell me about yourself", "introduce yourself", "how do you work",
    }
    if q in GREETING_EXACT: return "GENERAL"
    if any(q.startswith(g) for g in GREETING_STARTERS): return "GENERAL"
    if q in SMALL_TALK_EXACT or any(q.startswith(t) for t in SMALL_TALK_EXACT): return "GENERAL"
    if len(q) <= 10 and not any(w in q for w in ["doc", "pdf", "file", "web", "news", "cv"]):
        return "GENERAL"

    if len(documents_db) > 0:
        VAGUE_DOC = [
            "this doc", "my doc", "the doc", "this pdf", "my pdf", "the pdf",
            "this file", "my file", "the file", "this paper", "my paper", "the paper",
            "this report", "my report", "the report", "this document", "my document", "the document",
            "this assignment", "my assignment", "the assignment", "this text", "this article",
            "this cv", "my cv", "the cv", "this resume", "my resume",
            "in it", "about it", "from it", "summarize the", "summary of the",
            "from the document", "in the document", "from my pdf", "in my pdf",
            "tell me about this", "what this", "what is this", "about this",
        ]
        if any(t in q for t in VAGUE_DOC): return "DOCUMENT"

        for data in documents_db.values():
            fname      = data.get("filename", "")
            fname_stem = fname.rsplit(".", 1)[0].lower()
            if (fname_stem and len(fname_stem) > 3 and fname_stem in q) or fname.lower() in q:
                return "DOCUMENT"

        CV_WORDS = [
            "certification", "certifications", "certificate", "skills", "experience",
            "education", "qualification", "work history", "employment", "projects",
            "achievements", "objective", "summary", "profile", "contact",
        ]
        cv_uploaded = any(
            any(kw in data.get("filename", "").lower() for kw in ["cv", "resume", "curriculum"])
            for data in documents_db.values()
        )
        if cv_uploaded and any(w in q for w in CV_WORDS): return "DOCUMENT"

    WEB_TRIGGERS = [
        "current weather", "weather today", "weather of ", "weather in ",
        "temperature in ", "temperature of ", "temperature today",
        "today's temperature", "today's weather", "today's forecast", "forecast for ", "climate in ",
        "latest news", "breaking news", "news today", "today's news",
        "current events", "what happened today", "recently announced",
        "current price", "live score", "stock price", "price of bitcoin", "price of gold",
        "exchange rate", "petrol price", "fuel price", "oil price",
        "who won the match", "match result",
        "do a search", "search for", "look up", "find out the current",
        "current president", "current prime minister", "current pm", "current cm",
        "current chief minister", "current governor", "current ceo", "current chairman",
        "current minister", "current leader", "current head",
        "current top ", "current list", "current ranking",
        "top 5 ", "top 10 ", "top 3 ",
        "current gdp", "gdp of ", "current inflation", "inflation rate of ",
        "current unemployment", "economic growth of ",
        "universities in ", "colleges in ", "list of universities", "list of colleges",
        "population of ", "population in ",
    ]
    if any(t in q for t in WEB_TRIGGERS): return "WEB"

    import re as _re
    has_year = bool(_re.search(r'\b(19|20)\d{2}\b', q))
    FACTUAL_CONTEXT = [
        "pakistan", "india", "china", "usa", "uk", "iran", "turkey",
        "france", "germany", "italy", "russia", "brazil", "canada",
        "bangladesh", "afghanistan", "saudi arabia", "uae", "dubai",
        "lahore", "karachi", "islamabad", "delhi", "beijing",
        "population", "gdp", "inflation", "unemployment", "growth rate",
        "president", "prime minister", "chief minister", "governor",
        "election", "parliament", "government", "policy",
        "price", "rate", "index", "ranking", "score",
        "nuclear", "atomic", "warhead", "missile", "military",
        "army", "navy", "airforce", "troops", "defense",
        "countries with", "nations with", "states with",
    ]
    has_factual_context = any(t in q for t in FACTUAL_CONTEXT)

    KNOWLEDGE_STARTERS = (
        "what is ", "what was ", "what were ", "what are ",
        "how does ", "how do ", "how is ", "how are ",
        "why is ", "why are ", "why does ", "why do ",
        "explain ", "define ", "describe ",
        "can you explain", "can you describe",
        "difference between", "compare ", "what's the difference", "pros and cons",
        "who is ", "who was ", "who are ",
    )
    if any(q.startswith(s) for s in KNOWLEDGE_STARTERS):
        if has_year or has_factual_context:
            if len(documents_db) > 0:
                DOC_CONTEXT = [
                    "the paper", "this paper", "the doc", "this doc", "the pdf", "this pdf",
                    "the report", "the document", "the assignment", "the article", "the study",
                    "the cv", "this cv", "the resume", "uploaded", "you have", "i gave",
                    "in it", "about it",
                ]
                if any(t in q for t in DOC_CONTEXT): return "DOCUMENT"
            return "WEB"
        if len(documents_db) > 0:
            DOC_CONTEXT = [
                "the paper", "this paper", "the doc", "this doc", "the pdf", "this pdf",
                "the report", "the document", "the assignment", "the article", "the study",
                "the cv", "this cv", "the resume", "uploaded", "you have", "i gave",
            ]
            if any(t in q for t in DOC_CONTEXT): return "DOCUMENT"
        return "GENERAL"

    if len(documents_db) > 0:
        ACTION_ON_DOC = [
            "summarize", "summary", "overview", "main points", "key points", "features",
            "highlights", "conclusion", "introduction", "abstract", "findings", "results",
            "methodology", "tell me more", "more details", "in short", "briefly",
            "in brief", "tldr", "tl;dr",
            "skills", "experience", "certifications", "education",
            "projects", "achievements", "work history",
        ]
        if any(t in q for t in ACTION_ON_DOC): return "DOCUMENT"

    return "GENERAL"


# ─── Document processing ──────────────────────────────────────────────────────

def process_pdf(path: str) -> str:
    pages = []
    try:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(path)
        for i, page in enumerate(loader.lazy_load()):
            if i >= MAX_PAGES: break
            if page.page_content.strip():
                pages.append((i + 1, page.page_content))
    except Exception as e:
        logger.warning(f"PyPDFLoader failed: {e}")
    if not pages:
        try:
            import pymupdf
            doc = pymupdf.open(path)
            for i in range(min(len(doc), MAX_PAGES)):
                text = doc[i].get_text()
                if text.strip(): pages.append((i + 1, text))
            doc.close()
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}")
    return "\n\n".join(f"[Page {n}]\n{t}" for n, t in pages)


def process_docx(path: str) -> str:
    from docx import Document as DocxDocument
    return "\n".join(p.text for p in DocxDocument(path).paragraphs if p.text.strip())


def process_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_page_number(text: str) -> int:
    import re
    m = re.search(r"\[Page (\d+)\]", text)
    return int(m.group(1)) if m else 0


def strip_page_markers(text: str) -> str:
    import re
    return re.sub(r"\[Page \d+\]\n?", "", text).strip()


def format_docs(docs: List[Document]) -> str:
    return "\n\n".join(d.page_content for d in docs)


# ─── Persistence ──────────────────────────────────────────────────────────────

def save_documents_meta():
    slim = {
        doc_id: {
            "id": d["id"], "filename": d["filename"], "file_type": d["file_type"],
            "uploaded_at": d["uploaded_at"], "file_path": d["file_path"],
            "text_length": len(d.get("text_content", "")),
        }
        for doc_id, d in documents_db.items()
    }
    with open(DOCUMENTS_META_FILE, "w") as f:
        json.dump(slim, f, indent=2)


def load_documents_meta():
    global documents_db
    if not os.path.exists(DOCUMENTS_META_FILE): return
    try:
        with open(DOCUMENTS_META_FILE) as f:
            data = json.load(f)
        for doc_id, meta in data.items():
            if os.path.exists(meta.get("file_path", "")):
                documents_db[doc_id] = {**meta, "text_content": ""}
                doc_status[doc_id] = "ready"
    except Exception as e:
        logger.warning(f"Could not load documents meta: {e}")


def save_chats():
    with open(CHAT_HISTORY_FILE, "w") as f:
        json.dump({"chats": list(chats_db.values())}, f, indent=2, default=str)


def load_chats():
    global chats_db
    if not os.path.exists(CHAT_HISTORY_FILE): return
    try:
        with open(CHAT_HISTORY_FILE) as f:
            chats_db = {c["id"]: c for c in json.load(f).get("chats", [])}
    except Exception as e:
        logger.warning(f"Could not load chats: {e}")


# ─── FAISS indexing ───────────────────────────────────────────────────────────

def process_document_background(doc_id: str):
    global _vectorstore, _retriever
    with _processing_lock:
        doc_status[doc_id] = "indexing"
        try:
            data = documents_db.get(doc_id)
            if not data: doc_status[doc_id] = "error"; return

            processors = {"pdf": process_pdf, "docx": process_docx, "txt": process_txt}
            text = processors[data["file_type"]](data["file_path"])
            documents_db[doc_id]["text_content"] = text
            if not text.strip(): doc_status[doc_id] = "error"; return

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". "],
            )
            chunks = splitter.split_text(text)[:MAX_CHUNKS]
            docs = [
                Document(
                    page_content=strip_page_markers(c),
                    metadata={
                        "doc_id": doc_id, "doc_name": data["filename"],
                        "chunk": i,
                        "page": extract_page_number(c) or (i // 5) + 1,
                    },
                )
                for i, c in enumerate(chunks)
            ]

            emb = get_embeddings()
            index = None
            for i in range(0, len(docs), EMBEDDING_BATCH_SIZE):
                batch   = docs[i:i + EMBEDDING_BATCH_SIZE]
                new_idx = FAISS.from_documents(batch, emb)
                index   = new_idx if index is None else (index.merge_from(new_idx) or index)
                time.sleep(0.1)

            if index is not None:
                if _vectorstore is None: _vectorstore = index
                else: _vectorstore.merge_from(index)
                _vectorstore.save_local(INDEX_DIR)
                _retriever = _vectorstore.as_retriever(search_kwargs={"k": TOP_K})

            doc_status[doc_id] = "ready"
            save_documents_meta()
            logger.info(f"{doc_id} indexed OK")
        except Exception as e:
            logger.error(f"Indexing failed for {doc_id}: {e}")
            doc_status[doc_id] = "error"


def rebuild_index():
    global _vectorstore, _retriever
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, separators=["\n\n", "\n", ". "]
    )
    all_docs = []
    for doc_id, data in documents_db.items():
        text = data.get("text_content", "")
        if not text.strip():
            try:
                processors = {"pdf": process_pdf, "docx": process_docx, "txt": process_txt}
                text = processors[data["file_type"]](data["file_path"])
                documents_db[doc_id]["text_content"] = text
            except Exception: continue
        for i, c in enumerate(splitter.split_text(text)[:MAX_CHUNKS]):
            all_docs.append(Document(
                page_content=strip_page_markers(c),
                metadata={"doc_id": doc_id, "doc_name": data["filename"], "chunk": i},
            ))
    if not all_docs: _vectorstore = None; _retriever = None; return
    emb = get_embeddings()
    index = None
    for i in range(0, len(all_docs), EMBEDDING_BATCH_SIZE):
        batch   = all_docs[i:i + EMBEDDING_BATCH_SIZE]
        new_idx = FAISS.from_documents(batch, emb)
        index   = new_idx if index is None else (index.merge_from(new_idx) or index)
    _vectorstore = index
    _vectorstore.save_local(INDEX_DIR)
    _retriever = _vectorstore.as_retriever(search_kwargs={"k": TOP_K})


# ─── Lifespan ─────────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    load_chats()
    load_documents_meta()
    logger.info("Pre-loading embeddings model...")
    get_embeddings()
    global _vectorstore, _retriever
    idx_path = os.path.join(INDEX_DIR, "index.faiss")
    if os.path.exists(idx_path):
        try:
            _vectorstore = FAISS.load_local(INDEX_DIR, get_embeddings(), allow_dangerous_deserialization=True)
            _retriever   = _vectorstore.as_retriever(search_kwargs={"k": TOP_K})
            logger.info("FAISS index loaded")
        except Exception as e:
            logger.warning(f"Could not load FAISS index: {e}")
    yield
    save_chats()
    save_documents_meta()


# ─── App ──────────────────────────────────────────────────────────────────────

app = FastAPI(title="Hybrid RAG System", version="3.0", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)


# ─── Models ───────────────────────────────────────────────────────────────────

class SourceDoc(BaseModel):
    content: str; document_name: str; chunk_id: int; page: int = 0

class QueryRequest(BaseModel):
    question: str = Field(..., min_length=1, max_length=1000)
    chat_id:  Optional[str] = None

class QueryResponse(BaseModel):
    answer: str; sources: List[SourceDoc]; question: str
    chat_id: Optional[str] = None; source_type: str = "general"; route: str = "general"

class DocResponse(BaseModel):
    id: str; filename: str; file_type: str; uploaded_at: str
    text_length: int; status: str = "ready"

class DocStatusResponse(BaseModel):
    id: str; status: str

class ChatResp(BaseModel):
    id: str; title: str; created_at: str; messages: List[dict]


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _save_to_chat(chat_id, question, answer, sources, route="general"):
    if chat_id and chat_id in chats_db:
        ts = datetime.now().isoformat()
        chats_db[chat_id]["messages"].extend([
            {"role": "user",      "content": question, "sources": None,    "timestamp": ts, "route": "user"},
            {"role": "assistant", "content": answer,   "sources": sources, "timestamp": ts, "route": route},
        ])
        if chats_db[chat_id].get("title", "New Chat") == "New Chat":
            chats_db[chat_id]["title"] = question[:50]
        save_chats()


def _retrieve_docs(question: str, chat_id: str = None):
    global _active_doc
    if _vectorstore is None:
        pending = [d for d, s in doc_status.items() if s in ("pending", "indexing")]
        return (None, "⏳ Still indexing. Please wait.") if pending else \
               (None, "No documents indexed. Please upload a document first.")

    q_lower = question.lower()
    target_doc_id, target_doc_name = None, None

    for doc_id, data in documents_db.items():
        fname      = data.get("filename", "")
        fname_stem = fname.rsplit(".", 1)[0].lower()
        if (fname_stem and len(fname_stem) > 3 and fname_stem in q_lower) or fname.lower() in q_lower:
            target_doc_id, target_doc_name = doc_id, fname; break

    VAGUE_REFS = [
        "this doc", "my doc", "the doc", "this pdf", "my pdf", "the pdf",
        "this file", "my file", "this document", "the document",
        "this paper", "my paper", "the paper", "this report", "the report",
        "this assignment", "this cv", "my cv", "the cv", "this resume", "my resume",
        # Broader "this" references without explicit doc word
        "what this", "what is this", "this is about", "about this",
        "tell me about this", "explain this",
    ]
    if target_doc_id is None and any(t in q_lower for t in VAGUE_REFS):
        if len(documents_db) == 1:
            target_doc_id   = list(documents_db.keys())[0]
            target_doc_name = list(documents_db.values())[0].get("filename", "")
        elif len(documents_db) > 1:
            sd = sorted(documents_db.items(), key=lambda x: x[1].get("uploaded_at", ""), reverse=True)
            target_doc_id, target_doc_name = sd[0][0], sd[0][1].get("filename", "")

    if target_doc_id is None and chat_id and chat_id in _active_doc:
        sid = _active_doc[chat_id]
        if sid in documents_db:
            target_doc_id, target_doc_name = sid, documents_db[sid].get("filename", "")

    if target_doc_id is None and len(documents_db) == 1:
        target_doc_id   = list(documents_db.keys())[0]
        target_doc_name = list(documents_db.values())[0].get("filename", "")

    if target_doc_id and chat_id:
        _active_doc[chat_id] = target_doc_id

    broad_triggers = [
        "about", "summarize", "summary", "overview", "describe", "explain",
        "tell me", "give me", "introduction", "intro", "topic", "contents",
        "main", "key points", "features", "analysis", "detail", "in detail",
        "in short", "briefly", "skills", "experience", "certifications",
        "education", "findings", "results", "methodology", "conclusion",
    ]
    is_broad = any(t in q_lower for t in broad_triggers)
    want_k   = TOP_K_BROAD if is_broad else TOP_K

    # Cap fetch_k to actual index size to prevent FAISS errors on small indexes
    total_vectors = getattr(getattr(_vectorstore, "index", None), "ntotal", 0) or 100
    fetch_k = min(50 if target_doc_id else want_k * 2, total_vectors)
    if fetch_k == 0:
        return None, "No documents have been indexed yet. Please upload a document first."

    results = _vectorstore.similarity_search_with_relevance_scores(question, k=fetch_k)

    if target_doc_id:
        filtered = [(d, s) for d, s in results if d.metadata.get("doc_id") == target_doc_id]
        if not filtered:
            try:
                ntotal   = _vectorstore.index.ntotal
                all_r    = _vectorstore.similarity_search_with_relevance_scores(question, k=ntotal)
                filtered = [(d, s) for d, s in all_r if d.metadata.get("doc_id") == target_doc_id]
            except Exception:
                filtered = []
        if filtered:
            results = filtered
        else:
            st   = doc_status.get(target_doc_id, "unknown")
            name = target_doc_name or "the document"
            return (None, f"⏳ '{name}' is still being indexed. Please wait and try again.") \
                if st in ("pending", "indexing") \
                else (None, f"No content found for '{name}'. Try re-uploading the file.")

    good = [(d, s) for d, s in results if s >= RELEVANCE_THRESHOLD]
    if not good:
        if is_broad and results:                 good = results[:want_k]
        elif results and results[0][1] >= 0.20:  good = results[:2]
        else:
            return None, "I could not find relevant information in your uploaded documents for this question."

    good = good[:want_k]
    docs = [d for d, _ in good]
    return docs, [
        SourceDoc(content=d.page_content, document_name=d.metadata.get("doc_name", "Unknown"),
                  chunk_id=d.metadata.get("chunk", 0), page=d.metadata.get("page", 0))
        for d in docs
    ]


def _build_doc_prompt(question, docs, is_broad):
    context   = format_docs(docs)[:MAX_CONTEXT_CHARS]
    unique    = {d.metadata.get("doc_name", "") for d in docs}
    doc_label = f'"{next(iter(unique))}"' if len(unique) == 1 else "the uploaded documents"
    base = (
        f"IMPORTANT: Answer ONLY from the document excerpts below. "
        f"Do NOT use outside knowledge. Do NOT hallucinate. "
        f"{'The excerpts are from ' + doc_label + '.' if is_broad else 'If the answer is not in the excerpts, say so clearly.'}\n\n"
        f"---BEGIN DOCUMENT EXCERPTS---\n{context}\n---END DOCUMENT EXCERPTS---\n\n"
        f"Question: {question}\nAnswer (based strictly on the excerpts above):"
    )
    return base


def answer_from_documents(question, chat_id=None):
    docs, src_or_err = _retrieve_docs(question, chat_id=chat_id)
    if docs is None: return src_or_err, []
    q_lower  = question.lower()
    is_broad = any(t in q_lower for t in [
        "about", "summarize", "summary", "overview", "describe", "explain",
        "tell me", "give me", "introduction", "topic", "main", "key points",
        "features", "analysis", "improvements", "detail",
    ])
    resp = get_llm().invoke(_build_doc_prompt(question, docs, is_broad))
    return (resp.content if hasattr(resp, "content") else str(resp)), src_or_err


def answer_general(question):
    prompt = (
        "You are Hybrid RAG, a knowledgeable AI assistant. "
        "Answer directly without greeting or preamble. "
        "For technical questions explain clearly with key ideas and examples. "
        "For greetings respond warmly in one sentence. "
        "Never start with 'Hello', 'Hi', or 'Hey'.\n\n"
        f"Question: {question}\nAnswer:"
    )
    resp = get_llm().invoke(prompt)
    return resp.content if hasattr(resp, "content") else str(resp)


def answer_web(question):
    try:
        results = get_web_search().run(question)
        trimmed = results[:2000]
        prompt  = (
            "You are a helpful assistant. Answer accurately using these web results. Be concise.\n\n"
            f"Web Results:\n{trimmed}\n\nQuestion: {question}\nAnswer:"
        )
    except Exception as e:
        logger.warning(f"Web search failed: {e}")
        return answer_general(question)
    resp = get_llm().invoke(prompt)
    return resp.content if hasattr(resp, "content") else str(resp)


async def _stream_llm(prompt):
    async for chunk in get_llm().astream(prompt):
        text = chunk.content if hasattr(chunk, "content") else str(chunk)
        if text: yield text


async def _stream_doc_answer(question, docs):
    q_lower  = question.lower()
    is_broad = any(t in q_lower for t in [
        "about", "summarize", "summary", "overview", "describe", "explain",
        "tell me", "give me", "introduction", "topic", "main", "key points",
        "features", "analysis", "improvements", "detail",
    ])
    async for token in _stream_llm(_build_doc_prompt(question, docs, is_broad)):
        yield token


# ─── Endpoints ────────────────────────────────────────────────────────────────

ROUTE_MAP = {"DOCUMENT": "documents", "WEB": "web", "GENERAL": "general"}


@app.get("/api/health")
async def health():
    return {"status": "ok", "docs": len(documents_db), "chats": len(chats_db),
            "indexed": _vectorstore is not None}


@app.post("/api/query/stream")
async def query_stream(req: QueryRequest):
    loop           = asyncio.get_event_loop()
    route          = await loop.run_in_executor(None, classify_query, req.question)
    frontend_route = ROUTE_MAP.get(route, "general")
    logger.info(f"Stream: {route} | '{req.question[:60]}'")

    async def generate():
        yield _json.dumps({"type": "route", "route": frontend_route}) + "\n"
        sources, full_answer = [], []
        try:
            if route == "DOCUMENT":
                docs, src_or_err = await loop.run_in_executor(
                    None, lambda: _retrieve_docs(req.question, chat_id=req.chat_id)
                )
                if docs is None:
                    yield _json.dumps({"type": "token", "text": src_or_err}) + "\n"
                    full_answer.append(src_or_err)
                else:
                    sources = src_or_err
                    async for token in _stream_doc_answer(req.question, docs):
                        yield _json.dumps({"type": "token", "text": token}) + "\n"
                        full_answer.append(token)

            elif route == "WEB":
                try:
                    web_results = await loop.run_in_executor(None, lambda: get_web_search().run(req.question))
                    trimmed = web_results[:2000]
                    prompt  = (
                        "Answer accurately using these web results. Be concise.\n\n"
                        f"Web Results:\n{trimmed}\n\nQuestion: {req.question}\nAnswer:"
                    )
                except Exception:
                    prompt = (
                        "You are Hybrid RAG. Answer directly without greeting.\n\n"
                        f"Question: {req.question}\nAnswer:"
                    )
                async for token in _stream_llm(prompt):
                    yield _json.dumps({"type": "token", "text": token}) + "\n"
                    full_answer.append(token)

            else:
                prompt = (
                    "You are Hybrid RAG, a knowledgeable AI assistant. "
                    "Answer directly without greeting or preamble. "
                    "For technical questions explain clearly. "
                    "For greetings respond warmly in one sentence. "
                    "Never start with 'Hello', 'Hi', or 'Hey'.\n\n"
                    f"Question: {req.question}\nAnswer:"
                )
                async for token in _stream_llm(prompt):
                    yield _json.dumps({"type": "token", "text": token}) + "\n"
                    full_answer.append(token)

        except Exception as e:
            logger.error(f"Stream error: {e}")
            err = f"Error: {e}"
            yield _json.dumps({"type": "token", "text": err}) + "\n"
            full_answer.append(err)

        _save_to_chat(req.chat_id, req.question, "".join(full_answer),
                      [s.model_dump() for s in sources], route=frontend_route)
        yield _json.dumps({"type": "done", "sources": [s.model_dump() for s in sources],
                           "source_type": frontend_route}) + "\n"

    return StreamingResponse(generate(), media_type="application/x-ndjson")


@app.post("/api/query", response_model=QueryResponse)
async def query(req: QueryRequest):
    loop        = asyncio.get_event_loop()
    route       = await loop.run_in_executor(None, classify_query, req.question)
    source_type = ROUTE_MAP.get(route, "general")
    sources     = []
    if route == "DOCUMENT":
        answer, sources = await loop.run_in_executor(None, answer_from_documents, req.question, req.chat_id)
    elif route == "WEB":
        answer = await loop.run_in_executor(None, answer_web, req.question)
    else:
        answer = await loop.run_in_executor(None, answer_general, req.question)
    _save_to_chat(req.chat_id, req.question, answer, [s.model_dump() for s in sources], route=source_type)
    return QueryResponse(answer=answer, sources=sources, question=req.question,
                         chat_id=req.chat_id, source_type=source_type, route=source_type)


@app.post("/api/documents/upload", response_model=DocResponse)
async def upload(file: UploadFile = File(...)):
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ["pdf", "docx", "txt"]:
        raise HTTPException(400, f"Unsupported: .{ext}")
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, "File too large. Max 10 MB.")
    if ext == "pdf":
        try:
            import pymupdf
            doc = pymupdf.open(stream=content, filetype="pdf")
            pages = len(doc); doc.close()
            if pages > MAX_PAGES: raise HTTPException(413, f"PDF too long ({pages} pages, max {MAX_PAGES}).")
        except HTTPException: raise
        except Exception: pass
    doc_id = str(uuid.uuid4())
    path   = os.path.join(STORAGE_DIR, f"{doc_id}.{ext}")
    try:
        with open(path, "wb") as f: f.write(content)
        documents_db[doc_id] = {
            "id": doc_id, "filename": file.filename, "file_type": ext,
            "uploaded_at": datetime.now().isoformat(), "text_content": "", "file_path": path,
        }
        doc_status[doc_id] = "pending"
        asyncio.get_event_loop().run_in_executor(None, process_document_background, doc_id)
        return DocResponse(id=doc_id, filename=file.filename, file_type=ext,
                           uploaded_at=documents_db[doc_id]["uploaded_at"], text_length=0, status="pending")
    except HTTPException: raise
    except Exception as e:
        if os.path.exists(path): os.remove(path)
        raise HTTPException(500, str(e))


@app.get("/api/documents", response_model=List[DocResponse])
async def list_docs():
    return [DocResponse(id=d["id"], filename=d["filename"], file_type=d["file_type"],
                        uploaded_at=d["uploaded_at"], text_length=len(d.get("text_content", "")),
                        status=doc_status.get(d["id"], "ready"))
            for d in documents_db.values()]


@app.get("/api/documents/{doc_id}/status", response_model=DocStatusResponse)
async def get_doc_status(doc_id: str):
    if doc_id not in documents_db: raise HTTPException(404, "Not found")
    return DocStatusResponse(id=doc_id, status=doc_status.get(doc_id, "ready"))


@app.delete("/api/documents/{doc_id}")
async def delete_doc(doc_id: str):
    if doc_id not in documents_db: raise HTTPException(404, "Not found")
    path = documents_db[doc_id].get("file_path")
    if path and os.path.exists(path): os.remove(path)
    del documents_db[doc_id]
    if doc_id in doc_status: del doc_status[doc_id]
    asyncio.get_event_loop().run_in_executor(None, rebuild_index)
    save_documents_meta()
    return {"message": "Deleted"}


@app.get("/api/chats", response_model=List[ChatResp])
async def list_chats():
    return [ChatResp(id=c["id"], title=c.get("title", "New Chat"),
                     created_at=c.get("created_at", ""), messages=c.get("messages", []))
            for c in sorted(chats_db.values(), key=lambda x: x.get("created_at", ""), reverse=True)]


@app.post("/api/chats", response_model=ChatResp)
async def create_chat():
    cid = str(uuid.uuid4())
    chats_db[cid] = {"id": cid, "title": "New Chat", "created_at": datetime.now().isoformat(), "messages": []}
    save_chats()
    return ChatResp(**chats_db[cid])


@app.get("/api/chats/{chat_id}", response_model=ChatResp)
async def get_chat(chat_id: str):
    if chat_id not in chats_db: raise HTTPException(404, "Not found")
    return ChatResp(**chats_db[chat_id])


@app.patch("/api/chats/{chat_id}/title")
async def update_chat_title(chat_id: str, req: dict):
    if chat_id not in chats_db: raise HTTPException(404, "Not found")
    if "title" not in req:      raise HTTPException(400, "Title required")
    chats_db[chat_id]["title"] = req["title"]
    save_chats()
    return {"id": chat_id, "title": req["title"]}


@app.delete("/api/chats/{chat_id}")
async def delete_chat(chat_id: str):
    if chat_id not in chats_db: raise HTTPException(404, "Not found")
    del chats_db[chat_id]
    save_chats()
    return {"message": "Deleted"}


# ─── Serve React frontend (MUST be last) ─────────────────────────────────────
_static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
if os.path.exists(_static_dir):
    app.mount("/", StaticFiles(directory=_static_dir, html=True), name="static")
    logger.info(f"Serving frontend from {_static_dir}")
else:
    logger.info("No static/ dir — running in API-only / dev mode")

# ─── Entry point ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)))